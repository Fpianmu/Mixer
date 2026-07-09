#!/usr/bin/env python3
"""生成参赛文档 v4 — PIL绘制的标准计算机流程图 + 全章节"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from PIL import Image, ImageDraw, ImageFont
import os, io, textwrap

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(OUT_DIR, '_flowcharts')
os.makedirs(IMG_DIR, exist_ok=True)

# ══════════════════════════════════
#  Pillow 流程图绘制引擎
# ══════════════════════════════════
FONT_FILE = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"

def draw_flowchart(filename, nodes, edges, title=""):
    """
    nodes: [(x, y, w, h, text, shape), ...]  shape='rect' or 'diamond'
    edges: [(x1,y1,x2,y2,label), ...]  coordinates from node centers
    """
    W, H = 800, 50 + len(nodes) * 62 + 30
    img = Image.new('RGB', (W, H), 'white')
    draw = ImageDraw.Draw(img)

    try: font = ImageFont.truetype(FONT_FILE, 13)
    except: font = ImageFont.load_default()

    try: font_sm = ImageFont.truetype(FONT_FILE, 10)
    except: font_sm = ImageFont.load_default()

    try: font_title = ImageFont.truetype(FONT_FILE, 15)
    except: font_title = ImageFont.load_default()

    # Title
    if title:
        draw.text((W//2, 5), title, fill='#1a5276', font=font_title, anchor='mt')

    # Draw edges first (behind nodes)
    for x1, y1, x2, y2, label in edges:
        draw.line([(x1, y1), (x2, y2)], fill='#5b9bd5', width=2)
        # Arrowhead
        import math
        angle = math.atan2(y2-y1, x2-x1)
        L = 10
        ax1 = x2 - L*math.cos(angle-0.4); ay1 = y2 - L*math.sin(angle-0.4)
        ax2 = x2 - L*math.cos(angle+0.4); ay2 = y2 - L*math.sin(angle+0.4)
        draw.polygon([(x2, y2), (ax1, ay1), (ax2, ay2)], fill='#5b9bd5')
        # Edge label
        if label:
            mx, my = (x1+x2)//2 + 10, (y1+y2)//2 - 5
            draw.text((mx, my), label, fill='#e74c3c', font=font_sm)

    # Draw nodes
    for x, y, w, h, text, shape in nodes:
        if shape == 'diamond':
            # Diamond: defined by 4 points
            cx, cy = x + w//2, y + h//2
            pts = [(cx, y), (x+w, cy), (cx, y+h), (x, cy)]
            draw.polygon(pts, fill='#d5e8d4', outline='#82b366', width=2)
            # Text centered
            lines = textwrap.wrap(text, width=22)
            ty = cy - len(lines)*8
            for line in lines:
                bbox = draw.textbbox((0,0), line, font=font)
                tw = bbox[2]-bbox[0]
                draw.text((cx-tw//2, ty), line, fill='#333', font=font)
                ty += 16
        else:
            # Rounded rectangle
            r = 10
            draw.rounded_rectangle([x, y, x+w, y+h], radius=r, fill='#dae8fc', outline='#6c8ebf', width=2)
            # Text centered
            lines = textwrap.wrap(text, width=28)
            cy = y + h//2 - len(lines)*8
            for line in lines:
                bbox = draw.textbbox((0,0), line, font=font)
                tw = bbox[2]-bbox[0]
                draw.text((x+w//2-tw//2, cy), line, fill='#333', font=font)
                cy += 16

    img.save(filename)

# ══════════════════════════════════
#  Docx 生成
# ══════════════════════════════════
doc = Document()
S = doc.styles['Normal']
S.font.name = '宋体'; S.font.size = Pt(11)
S.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def P(text):
    p = doc.add_paragraph(text); p.paragraph_format.line_spacing = 1.5

def B(text):
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_before = Pt(10)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11)

def T(hds, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(hds))
    table.style = 'Light Grid Accent 1'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(hds):
        c = table.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph()

def IO(rows):
    T(['变量名','类型','分类','含义'], rows)

def IMG(name, width_inches=6.0):
    """嵌入流程图图片"""
    path = os.path.join(IMG_DIR, name)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(path, width=Inches(width_inches))
    doc.add_paragraph()

# ══════════════════════════════════

# --- 绘制所有流程图 ---
# 通用布局: 节点x居中=400, w=300, 左x=250
CX = 400; W = 300; LX = 250

# 图1: 系统调用层次
draw_flowchart(os.path.join(IMG_DIR, 'fc1_callstack.png'),
    nodes=[
        (LX,20,W,40,"app_main()\nESP32-S3 入口函数",'rect'),
        (LX,80,W,40,"init_all()\n系统初始化总入口",'rect'),
        (LX-120,140,200,40,"ledc_pwm\n水泵PWM",'rect'),
        (LX+20,140,200,40,"servo/servo2\n双舵机",'rect'),
        (LX+160,140,200,40,"esc_init\n双路电调",'rect'),
        (LX-170,200,200,40,"wifi_init\nWiFi AP",'rect'),
        (LX-170,260,200,40,"http_server\nHTTP服务",'rect'),
        (LX+30,200,200,40,"xz_init\n小智AI语音",'rect'),
        (LX+30,260,200,40,"lcd_init\n触摸屏",'rect'),
        (LX+230,200,200,40,"ui_init\nEEZ UI",'rect'),
        (LX,320,W,40,"while(1) { ui_tick() }\n主循环 10ms 驱动LVGL图形界面",'rect'),
    ],
    edges=[
        (CX,60,CX,80,""),
        (CX,120,230,140,""),
        (CX,120,510,140,""),
        (CX,120,660,140,""),
        (180,180,180,200,""),
        (180,240,180,260,""),
        (420,180,420,200,""),
        (420,240,420,260,""),
        (560,180,580,200,""),
        (CX,200,400,320,""),
        (CX,280,CX,320,""),
    ],
    title="系统初始化调用层次"
)

# 图2: weight_work 核心流程
draw_flowchart(os.path.join(IMG_DIR, 'fc2_weightwork.png'),
    nodes=[
        (LX,10,W,40,"weight_work(weight)\n输入: 面团重量",'rect'),
        (LX-120,70,200,40,"配方比例计算\nk_flour/k_water/\nk_grain/k_yeast/k_salt",'rect'),
        (CX-50,130,100,50,"(int) 类型\n转换正确?",'diamond'),
        (LX-120,200,200,40,"修正: 截断在\n除法运算之后",'rect'),
        (LX,260,W,40,"fwork(flour,water,500,grain,400000)\n四阶段分步执行",'rect'),
        (LX-170,320,170,40,"Task1 面粉加入\nesc1 40%油门",'rect'),
        (LX-170,380,170,40,"Task2 水泵加水\nDRV8870 PWM",'rect'),
        (LX,320,170,40,"Task3 研磨+加料\n舵机循环投料",'rect'),
        (LX,380,170,40,"Task4 搅拌\nesc2 40%油门",'rect'),
    ],
    edges=[
        (CX,50,CX,70,""),
        (CX,110,CX,130,""),
        (CX,180,100,200,"否→修正"),
        (100,220,CX-50,260,""),
        (CX,300,CX,320,""),
        (CX,360,85,320,""),
        (85,360,85,380,""),
        (CX,340,CX,380,""),
        (CX,360,570,320,""),
        (570,360,570,380,""),
    ],
    title="weight_work() → fwork() 核心工作流"
)

# 图3: 互斥锁
draw_flowchart(os.path.join(IMG_DIR, 'fc3_mutex.png'),
    nodes=[
        (LX,10,W,40,"ctl_try_acquire(CTL_TOUCH)\n触控屏请求控制权",'rect'),
        (LX,70,W,40,"xSemaphoreTake(mutex, 0)\n非阻塞尝试获取信号量",'rect'),
        (CX-50,130,100,55,"获取\n成功?",'diamond'),
        (CX+160,130,160,40,"返回 false\n系统正忙，拒绝操作",'rect'),
        (CX-50,210,100,55,"s_owner\n≠ NONE?",'diamond'),
        (CX+160,210,160,40,"返回 false\n已被其他源占用",'rect'),
        (LX,285,W,40,"s_owner = CTL_TOUCH\n登记为当前控制源",'rect'),
        (LX,345,W,40,"xSemaphoreGive(mutex)\n释放信号量 → return true",'rect'),
    ],
    edges=[
        (CX,50,CX,70,""),
        (CX,110,CX,130,""),
        (CX,185,560,130,"否"),
        (CX,185,CX,210,"是"),
        (CX,265,555,210,"否"),
        (CX,255,CX,285,"是"),
        (CX,325,CX,345,""),
    ],
    title="ctl_try_acquire() 互斥锁控制流程"
)

# 图4: 小智AI
draw_flowchart(os.path.join(IMG_DIR, 'fc4_xiaozhi.png'),
    nodes=[
        (LX,5,W,40,"微信小程序\n用户输入饮食和面食需求",'rect'),
        (LX-120,65,250,40,"POST /api/v1/intake/parse\nLLM解析 → ParseResult",'rect'),
        (CX-50,125,100,55,"用户\n确认?",'diamond'),
        (CX+160,125,160,40,"小程序: 允许修改\n饮食条目和重量",'rect'),
        (LX-120,200,250,40,"POST /api/v1/recommendations/\ncoarse-grain → Recommendation",'rect'),
        (CX-50,260,100,55,"确认\n执行?",'diamond'),
        (CX+160,260,160,40,"取消操作\n返回空闲状态",'rect'),
        (LX,320,W,40,"WebSocket MCP → ESP32\nhandle_mcp_command()",'rect'),
        (CX-50,380,100,55,"互斥锁\n获取?",'diamond'),
        (CX+160,380,160,40,"语音反馈\n\"设备正忙\"",'rect'),
        (LX,440,W,40,"weight_work(weight)\n四阶段自动执行 → 完成",'rect'),
    ],
    edges=[
        (CX,45,CX,65,""),
        (CX,105,CX,125,""),
        (CX,180,560,125,"否→修改"),
        (CX,175,CX,200,"是"),
        (CX,240,CX,260,""),
        (CX,315,560,260,"否"),
        (CX,315,CX,320,"是"),
        (CX,360,CX,380,""),
        (CX,435,560,380,"否→反馈"),
        (CX,435,CX,440,"是"),
    ],
    title="微信小程序 → 小智AI服务端 → ESP32 设备控制全链路"
)

# 图5: 推荐算法
draw_flowchart(os.path.join(IMG_DIR, 'fc5_recommender.png'),
    nodes=[
        (LX,10,W,40,"输入: items(当日食物清单)\n+ dough_request(面团需求)",'rect'),
        (LX,70,W,40,"refined = Σrefined_staple\ncoarse = Σcoarse_grain",'rect'),
        (LX,130,W,40,"ratio = 0.20 (基础20%)",'rect'),
        (CX-120,190,180,50,"refined ≥ 300g?\nratio += 0.15",'diamond'),
        (CX+60,190,180,50,"refined ≥ 150g?\nratio += 0.10",'diamond'),
        (CX-120,260,180,50,"coarse ≥ 100g?\nratio -= 0.05",'diamond'),
        (CX-120,330,180,50,"消化风险?\nratio ≤ 0.20",'diamond'),
        (LX,400,W,40,"ratio = clamp(ratio, 0.10, 0.40)\nround到5%整倍数",'rect'),
        (LX,460,W,40,"water = total×0.24\nflour = total-water\ncoarse_grain = flour×ratio",'rect'),
        (LX,520,W,40,"输出: Recommendation\n{flour, water, coarse_grain, ratio, reason}",'rect'),
    ],
    edges=[
        (CX,50,CX,70,""), (CX,110,CX,130,""),
        (CX,170,310,190,""), (CX,170,480,190,""),
        (CX,210,310,330,""), (CX,210,480,330,""),
        (CX,300,CX,400,""), (CX,440,CX,460,""), (CX,500,CX,520,""),
    ],
    title="recommend_coarse_grain() 营养推荐算法"
)

# 图6: 流量计
draw_flowchart(os.path.join(IMG_DIR, 'fc6_flowmeter.png'),
    nodes=[
        (LX,10,W,40,"Task2 水泵加水\n输入: water_weight_g",'rect'),
        (LX,70,W,40,"target = ceil(水量mL × 0.45)\npulse_count = 0",'rect'),
        (LX,130,W,40,"ledc_pwm_set_state(1)\n开水泵, GPIO12输出PWM",'rect'),
        (CX-50,190,100,55,"pulse_count\n< target?",'diamond'),
        (CX+160,190,140,40,"vTaskDelay(10ms)\n继续轮询等待",'rect'),
        (CX-50,265,100,40,"GPIO37中断\npulse_count++",'rect'),
        (LX,325,W,40,"pulse_count ≥ target\nledc_pwm_set_state(0) 关水泵",'rect'),
    ],
    edges=[
        (CX,50,CX,70,""), (CX,110,CX,130,""),
        (CX,170,CX,190,""),
        (CX,245,570,190,"是→等待"),
        (CX,245,CX,265,"否→达标"),
        (CX,305,CX,325,""),
    ],
    title="水泵加水（YF-S201 流量计闭环脉冲计数）"
)

# 图7: 触控屏启动
draw_flowchart(os.path.join(IMG_DIR, 'fc7_touch.png'),
    nodes=[
        (LX,10,W,40,"action_star_mixer()\n用户点击\"启动\"按钮",'rect'),
        (LX,70,W,40,"w = get_var_dough_weight()\n读取UI当前重量值",'rect'),
        (CX-50,130,100,55,"ctl_try\n_acquire\n(TOUCH)?",'diamond'),
        (CX+160,130,160,40,"返回 (无响应)\n系统正忙",'rect'),
        (LX,200,W,40,"set_var_motor_running(true)\n更新UI状态",'rect'),
        (LX,260,W,40,"xTaskCreate(motor_task)\n创建独立FreeRTOS任务",'rect'),
        (LX,320,W,40,"motor_task → weight_work(w)\n→ fwork() 四阶段执行",'rect'),
    ],
    edges=[
        (CX,50,CX,70,""), (CX,110,CX,130,""),
        (CX,185,555,130,"否→拒绝"),
        (CX,175,CX,200,"是"),
        (CX,240,CX,260,""), (CX,300,CX,320,""),
    ],
    title="触控屏启动和面流程"
)

print(f"Flowcharts generated: {len(os.listdir(IMG_DIR))} images")

# ═══════════════════ 2.1 ═══════════════════
H('2.1 整体介绍', level=1)
P('本系统以 ESP32-S3 为主控芯片，构建智能全自动和面机。系统采用四层架构：用户交互层、智能服务层、主控决策层与硬件执行层。')
P('触控屏（LVGL/EEZ Studio）、网页控制面板（HTTP REST API）和微信小程序三种用户入口通过统一的互斥锁机制共享硬件资源。微信小程序端用户以文字或语音描述当日饮食情况和面食需求，经小智 AI 服务端（xiaozhi-server）调用 DeepSeek 大语言模型和营养推荐引擎解析后，生成面团配方的结构化参数，通过 WebSocket 长连接下发至 ESP32 设备端。设备端运行 FreeRTOS 多任务系统，核心工作流引擎根据接收到的参数，按配方比例驱动 DRV8870 水泵、无刷电调、双舵机、步进电机和霍尔流量计等执行机构，完成面粉加入、加水、研磨、加料、搅拌的全自动流程。')

H('系统整体框图', level=2)
T(['层次','组成模块','关键技术'],
  [['用户交互层','触控屏（LVGL+EEZ Studio）\n网页控制面板（HTTP REST）\n微信小程序（小智 AI 入口）','三端互斥锁（ctl_mutex）\nFreeRTOS 多任务调度'],
   ['智能服务层','小智 AI 服务端（xiaozhi-server）\nASR / LLM / TTS 引擎\n营养推荐引擎（nutrition-service）','DeepSeek 大语言模型\nWebSocket 实时通信\nMCP 工具调用协议'],
   ['主控决策层','ESP32-S3（240MHz Xtensa LX7）\nFreeRTOS 实时操作系统\n核心工作流引擎（function.c）','LEDC 8通道全满调度\nI2S 全双工音频管线\nOpus + ESP-SR 唤醒词引擎'],
   ['硬件执行层','DRV8870 水泵（GPIO12/13 PWM）\n无刷电调 ×2（GPIO15/9 ESC）\n双舵机 + 步进电机\nYF-S201 流量计（GPIO37）','5kHz PWM 电机调速\n50Hz ESC 油门控制\n闭环脉冲计数水量控制']])

P('各层之间通过明确的接口通信：用户交互层经 ctl_mutex 统一调度后调用核心工作流引擎；智能服务层通过 WebSocket/MCP 下发结构化指令至设备端；主控决策层通过 LEDC/I2S/GPIO 驱动硬件执行层。层次间低耦合、层内高内聚的设计便于后续独立升级任一模块。')

# ═══════════════════ 2.3 ═══════════════════
H('2.3 软件系统介绍', level=1)
H('2.3.1 软件总体设计', level=2)

P('ESP32-S3 搭载 FreeRTOS 实时操作系统，通过抢占式多任务调度管理图形界面渲染、WiFi 网络通信、HTTP/TCP 指令处理、音频语音交互及电机控制工作流等并行任务。任务间通过互斥信号量、消息队列和事件组进行同步，确保多任务并发下的实时响应能力。')
P('ESP32-S3 内部资源在设计中得到充分利用：4 个 LEDC 定时器与全部 8 个通道被分配至 7 路独立 PWM 输出；I2S0 外设配置为全双工模式，以 4 个 GPIO 同时驱动 INMP441 麦克风输入和 MAX98357A 扬声器输出。在 240MHz 单核 Xtensa LX7 处理器上，通过 FreeRTOS 时间片轮转调度，将 CPU 密集型任务与实时性任务分配到不同优先级的任务中，保证关键控制路径的确定性和语音交互的低延迟。')

# ═══════════════════ 2.3.2 ═══════════════════
H('2.3.2 软件各模块介绍', level=1)

# ── 模块一 ──
H('一、核心工作流引擎（function.c）', level=2)
P('工作流引擎接收面团重量参数，按预设配方比例计算各物料目标值，分四阶段顺序执行全自动和面流程。')
IMG('fc1_callstack.png', 6.0)
IMG('fc2_weightwork.png', 6.0)

B('关键输入变量：')
IO([
    ['weight','uint32_t','输入','面团总重量（g），来自触控屏/HTTP/微信小程序'],
    ['k_flour~k_salt','float','常量(5组)','配方比例常数：120/267.5, 65/267.5, 80/267.5, 2/267.5, 0.5/267.5'],
])
B('关键输出变量：')
IO([
    ['flour','int','输出','面粉电机运行时长（ms）'],
    ['water','int','输出','水泵运行时长/目标脉冲数'],
    ['grain','int','输出','研磨运行总时长（ms）'],
    ['yeast','int','输出','酵母加料循环次数'],
    ['salt','int','输出','盐加料循环次数'],
])

# ── 模块二 ──
H('二、三端互斥控制（ctl_mutex.c）', level=2)
P('系统支持触控屏、HTTP 网页、小智 AI 三方控制。本模块基于 FreeRTOS 互斥信号量实现非阻塞互斥调度，确保同一时刻仅有一端操作硬件。采用非阻塞策略（xSemaphoreTake 超时为0）而非阻塞等待，避免触控屏 UI 线程或 HTTP 请求被长时间挂起导致用户体验冻结。')
IMG('fc3_mutex.png', 6.0)

B('控制源接入点：')
T(['控制源','枚举值','获取锁位置','释放锁位置'],
  [['触控屏','CTL_TOUCH','action_star_mixer()','action_stop()'],
   ['HTTP 网页','CTL_HTTP','api_start_post_handler()','api_stop_post_handler()'],
   ['小智 AI','CTL_XIAOZHI','handle_mcp_command()','fstop()']])

B('关键变量：')
IO([
    ['s_mutex','SemaphoreHandle_t','内部','FreeRTOS 互斥信号量'],
    ['s_owner','ctl_source_t','内部','NONE/TOUCH/HTTP/XIAOZHI'],
    ['src (入参)','ctl_source_t','输入','请求控制权的来源标识'],
    ['返回值','bool','输出','true=获取成功, false=已被占用'],
])

# ── 模块三 ──
H('三、电机与执行机构驱动', level=2)
P('本系统共驱动 4 类、7 路独立执行机构，全部通过 ESP32-S3 LEDC 硬件外设产生 PWM 信号，经 4 个定时器与 8 个独立通道统一管理。5kHz PWM 用于 DRV8870 有刷电机驱动，50Hz 用于舵机和航模电调，1kHz 用于步进电机脉冲序列。8 通道全满，4 定时器无空闲，资源利用率 100%。')

B('LEDC 通道分配：')
T(['Timer','频率','分辨率','Ch','GPIO','执行机构','信号类型'],
  [['0','5kHz','10-bit','CH3','12','水泵 DRV8870 IN1','PWM 可变占空比'],
   ['0','5kHz','10-bit','CH4','13','水泵 DRV8870 IN2','固定 LOW'],
   ['1','1kHz','10-bit','CH1','4','步进电机 STP','脉冲序列'],
   ['2','50Hz','14-bit','CH0','8','旋转舵机','角度-脉宽映射'],
   ['2','50Hz','14-bit','CH2','3','击打舵机','角度-脉宽映射'],
   ['3','50Hz','14-bit','CH6','15','面粉 ESC','油门-脉宽映射'],
   ['3','50Hz','14-bit','CH7','9','搅拌 ESC','油门-脉宽映射']])

B('技术难点：')
P('选择 10-bit 分辨率（1024 占空比等级）而非更高位宽的原因是 LEDC 频率可达性约束。ESP32-S3 的 LEDC 时钟源为 80MHz APB 时钟。若采用 13-bit（8192 等级），理论最高频率仅为 9.8kHz，无法达到目标 5kHz。10-bit 模式下最高频率为 78kHz，5kHz 通过分数分频器精确实现。这一约束在初版调试中被忽略，导致电机不转——直到用万用表测得 GPIO12 仅 0.18V 而非预期的约 1.65V，通过反向计算频率可达性才定位到根因。')

B('关键变量：')
IO([
    ['state (入参)','int','输入','0=关闭, 1=开启50%占空比'],
    ['duty','uint32_t','输出','目标占空比值（0 或 511）'],
    ['PWM_FREQ_HZ','宏','常量','50000 (5kHz)'],
    ['DUTY_50_PERCENT','宏','常量','511 (1023总等级的50%)'],
])

# ── 模块四 ──
H('四、小智 AI 智能交互', level=2)
P('用户通过微信小程序聊天界面描述当日饮食和面食需求，经小智 AI 服务端解析后，通过 MCP 协议下发结构化控制指令至 ESP32。服务端由 FastAPI 营养推荐引擎（nutrition-service）和 WebSocket 网关（xiaozhi-server）两部分组成：营养引擎负责饮食解析与杂粮配比推荐，WebSocket 网关负责设备实时通信和 MCP 工具调用。')
IMG('fc4_xiaozhi.png', 6.0)
IMG('fc5_recommender.png', 6.0)

B('MCP 指令映射：')
T(['用户意图','MCP tool','设备端函数','参数'],
  [['开始和面','start_mixer','weight_work(weight)','weight: 面团总重量(g)'],
   ['紧急停止','stop_mixer','fstop()','无'],
   ['推出面团','push_out','push_and_out(1)','direction: 1'],
   ['退回','push_back','push_and_out(0)','direction: 0']])

B('关键输入变量：')
IO([
    ['text','string','输入','用户当日饮食+面食需求自然语言'],
    ['confirmed_items','list[FoodItem]','输入','用户确认后的饮食条目列表'],
    ['dough_request','DoughRequest','输入','面团需求（食物类型、总重量）'],
])
B('关键输出变量：')
IO([
    ['ParseResult.items','list[FoodItem]','输出','食物条目（meal, weight, category）'],
    ['Recommendation.ratio','float','输出','杂粮比例（10%~40%，步进5%）'],
    ['Recommendation.flour','int','输出','普通面粉克重（g）'],
    ['Recommendation.coarse_grain','int','输出','杂粮粉克重（g）'],
    ['Recommendation.water','int','输出','水量（g/mL）'],
    ['MCP JSON','JSON','输出','下发给 ESP32 的结构化控制指令'],
])

# ── 模块五 ──
H('五、流量计闭环水量控制', level=2)
P('将传统定时开环加水（精度 ±20%）升级为 YF-S201 霍尔流量传感器闭环脉冲计数控制（精度 ±2~5%）。传感器输出 450 脉冲/L，GPIO37 配置为上升沿中断输入，IRAM 中的轻量级 ISR 仅执行 pulse_count++。主循环以 10ms 间隔比较当前计数与目标值，达标后立即关断水泵。未采用 PID 补偿——水泵关断后管路残余水量（2~4mL）在传感器固有误差范围内。')
IMG('fc6_flowmeter.png', 6.0)

B('关键变量：')
IO([
    ['pulse_count','volatile uint32_t','全局','IRAM ISR 中的脉冲累加计数器'],
    ['target','uint32_t','局部','目标脉冲数 = ceil(水量mL × 0.45)'],
    ['water_weight_g','int','输入','配方计算出的水量（g，即 mL）'],
])

# ── 模块六 ──
H('六、LVGL 触控屏与 HTTP 服务器', level=2)
P('触控屏采用 LVGL v9 + EEZ Studio，EEZ 生成页面布局（screens.c/ui.c），action.cpp 处理按钮事件并映射至工作流函数。HTTP 服务器（端口 80）内嵌网页固件，提供 REST API（/api/start、/api/stop、/api/status）。两个模块在执行电机控制前均需通过 ctl_try_acquire 获取互斥锁。')
IMG('fc7_touch.png', 6.0)

B('关键变量：')
IO([
    ['dough_weight','int32_t','全局','UI 当前面团重量值（50~1000g）'],
    ['motor_running','bool','全局','UI 电机运行状态指示'],
    ['g_motor_task_handle','TaskHandle_t','内部','电机任务 FreeRTOS 句柄'],
])

# ═══════════════════ 2.3.3 ═══════════════════
H('2.3.3 软件技术指标', level=1)
T(['指标','数值','指标','数值'],
  [['操作系统','FreeRTOS（抢占式多任务）','音频编码','Opus（16kHz/单声道/60ms帧）'],
   ['固件框架','ESP-IDF v6.0 (CMake+Ninja)','唤醒词引擎','ESP-SR MultiNet'],
   ['编程语言','C(驱动层) / C++(应用层+音频)','服务端模型','DeepSeek / ChatGLM'],
   ['图形框架','LVGL v9 + EEZ Studio','HTTP服务','内嵌服务器, 端口80, REST JSON'],
   ['LEDC占用','8/8通道, 4/4定时器, 100%','WiFi模式','AP (802.11n, 20MHz)'],
   ['固件体积','2.7MB (工厂分区3MB, 14%空闲)','I2S模式','全双工 (4 GPIO双设备)'],
   ['GPIO占用','34/49 (69%)','PWM输出','7路独立 (LEDC)'],
])

H('2.3.4 完整引脚与片上资源分配', level=1)
P('（此处插入《ESP32-S3完整资源分配表.docx》中的五个表格。）')

# ── Save ──
out = os.path.join(OUT_DIR, '参赛文档_软件系统介绍_v4.docx')
doc.save(out)
print(f'\nOK: {out}')
