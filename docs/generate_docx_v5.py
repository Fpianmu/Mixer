#!/usr/bin/env python3
"""v5 — 中文字体 + 宽松排版 + 标准流程图"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
import os, textwrap, math

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(OUT_DIR, '_flowcharts')
os.makedirs(IMG_DIR, exist_ok=True)

# 中文字体
FONT_PATH = "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc"
FONT_PATH_REG = "/usr/share/fonts/opentype/noto/NotoSansCJK-Medium.ttc"

def draw_fc(filename, title, steps, decision_steps, edges):
    """
    统一绘制一张流程图。
    steps: [(x, y, w, h, text, is_diamond)]
    edges: [(from_idx, to_idx, label)]
    自动计算画布大小
    """
    W = 1200
    max_y = max(y+h for (_,_,_,_,y,h,_,_) in steps) if 'steps' in dir() else 800
    # Recalculate from the actual data
    pass  # 下面用新逻辑

# ── 新流程图引擎 ──
def make_flowchart(filename, title, nodes, edges):
    """
    nodes: [(text, is_diamond)]  自动布局, 垂直排列
    edges: [(from_idx, to_idx, label)]
    """
    N = len(nodes)
    # 自动计算间距
    SPACING = 22  # 节点间距
    NODE_W, NODE_H = 580, 62    # 矩形
    DIAMOND_W, DIAMOND_H = 280, 80  # 菱形
    LEFT_X = 310  # 矩形左边界(居中)
    CX = NODE_W // 2 + LEFT_X  # 中心x = 310 + 290 = 600

    # 第一遍: 计算每个节点的 y
    positions = []
    cur_y = 50
    for i, (text, is_diamond) in enumerate(nodes):
        if is_diamond:
            h = DIAMOND_H
            x = CX - DIAMOND_W // 2
            w = DIAMOND_W
        else:
            h = NODE_H
            x = LEFT_X
            w = NODE_W
        positions.append((x, cur_y, w, h))
        cur_y += h + SPACING

    H = cur_y + 20
    W = 1200
    img = Image.new('RGB', (W, H), 'white')
    draw = ImageDraw.Draw(img)

    try:
        font = ImageFont.truetype(FONT_PATH, 17)
        font_sm = ImageFont.truetype(FONT_PATH_REG, 14)
        font_edge = ImageFont.truetype(FONT_PATH_REG, 12)
        font_title = ImageFont.truetype(FONT_PATH, 22)
    except:
        font = ImageFont.load_default()
        font_sm = font; font_edge = font; font_title = font

    # 先画边
    for fi, ti, label in edges:
        fx, fy, fw, fh = positions[fi]
        tx, ty, tw, th = positions[ti]
        # 从底中心到顶中心
        x1, y1 = fx + fw // 2, fy + fh
        x2, y2 = tx + tw // 2, ty

        # 如果有分支(菱形→侧边矩形),调整坐标
        if nodes[fi][1]:  # 菱形
            if tx > fx + fw // 2:  # 右侧分支
                x1, y1 = fx + fw, fy + fh // 2
                x2, y2 = tx, ty + th // 2
            elif tx < fx + fw // 2:  # 左侧分支
                x1, y1 = fx, fy + fh // 2
                x2, y2 = tx + tw, ty + th // 2

        draw.line([(x1, y1), (x2, y2)], fill='#5b9bd5', width=2)

        # 箭头
        angle = math.atan2(y2-y1, x2-x1)
        L = 14
        ax1 = x2 - L*math.cos(angle-0.45); ay1 = y2 - L*math.sin(angle-0.45)
        ax2 = x2 - L*math.cos(angle+0.45); ay2 = y2 - L*math.sin(angle+0.45)
        draw.polygon([(x2, y2), (ax1, ay1), (ax2, ay2)], fill='#5b9bd5')

        if label:
            mx, my = (x1+x2)//2 + 12, (y1+y2)//2 - 6
            draw.text((mx, my), label, fill='#c0392b', font=font_edge)

    # 再画节点
    for i, (text, is_diamond) in enumerate(nodes):
        x, y, w, h = positions[i]
        if is_diamond:
            cx, cy = x + w//2, y + h//2
            pts = [(cx, y), (x+w, cy), (cx, y+h), (x, cy)]
            draw.polygon(pts, fill='#d4edda', outline='#28a745', width=3)
            lines = textwrap.wrap(text, width=16)
            ty = cy - len(lines)*12
            for line in lines:
                bb = draw.textbbox((0,0), line, font=font_sm)
                tw = bb[2]-bb[0]
                draw.text((cx-tw//2, ty), line, fill='#222', font=font_sm)
                ty += 22
        else:
            r = 14
            draw.rounded_rectangle([x, y, x+w, y+h], radius=r, fill='#dae8fc', outline='#6c8ebf', width=3)
            lines = textwrap.wrap(text, width=32)
            cy = y + h//2 - len(lines)*12
            for line in lines:
                bb = draw.textbbox((0,0), line, font=font)
                tw = bb[2]-bb[0]
                draw.text((x+w//2-tw//2, cy), line, fill='#222', font=font)
                cy += 22

    # 标题
    draw.text((W//2, 12), title, fill='#1a5276', font=font_title, anchor='mt')

    img.save(filename)
    print(f"  {os.path.basename(filename)} ({W}x{H})")

# ═══════════════════ 绘制 7 张流程图 ═══════════════════

# 1: 系统调用层次
make_flowchart(os.path.join(IMG_DIR, 'fc1_stack.png'), '系统初始化调用层次',
    nodes=[
        ("app_main()\nESP32-S3 入口函数", False),
        ("init_all()\n系统初始化总入口", False),
        ("ledc_pwm / servo/servo2 / esc_init\n水泵PWM + 双舵机 + 双路电调", False),
        ("wifi_init / http_server / xz_init\nWiFi热点 + HTTP服务 + 小智AI语音", False),
        ("bsp_lcd_init / ui_init\nLCD液晶屏 + EEZ图形界面", False),
        ("while(1) { ui_tick() }\n主循环 10ms 驱动LVGL", False),
    ],
    edges=[(0,1,""),(1,2,""),(2,3,""),(3,4,""),(4,5,"")]
)

# 2: 核心工作流
make_flowchart(os.path.join(IMG_DIR, 'fc2_weightwork.png'), 'weight_work() → fwork() 核心工作流',
    nodes=[
        ("weight_work(weight)\n输入: 面团重量(来自触控屏/HTTP/小程序)", False),
        ("配方比例计算\nk_flour / k_water / k_grain / k_yeast / k_salt", False),
        ("int 强制类型转换\n截断置于除法运算之前？", True),
        ("fwork(flour, water, 500, grain, 400000)\n四阶段分步执行", False),
        ("Task1 面粉加入: ESC1 40%油门 + 磁铁阀GPIO18 ON → 延时关闭", False),
        ("Task2 水泵加水: DRV8870 PWM ON → 流量计脉冲达标 → OFF", False),
        ("Task3 研磨+加料: 舵机循环投料 + 研磨间歇ON/OFF", False),
        ("Task4 搅拌: ESC2 40%油门, 延时400s → 关闭 → IDLE", False),
    ],
    edges=[(0,1,""),(1,2,""),(2,4,"是→bug"),(2,3,"否→修正"),(3,4,""),(4,5,""),(5,6,""),(6,7,"")]
)

# 3: 互斥锁
make_flowchart(os.path.join(IMG_DIR, 'fc3_mutex.png'), 'ctl_try_acquire() 互斥锁控制流程',
    nodes=[
        ("ctl_try_acquire(CTL_TOUCH)\n触控屏请求控制权", False),
        ("xSemaphoreTake(mutex, 0)\n非阻塞尝试获取互斥信号量", False),
        ("信号量获取成功 ?", True),
        ("s_owner == NONE ?\n(当前无其他源占用)", True),
        ("s_owner = CTL_TOUCH\n登记为当前控制源", False),
        ("xSemaphoreGive(mutex)\n释放信号量 → 返回 true", False),
    ],
    edges=[(0,1,""),(1,2,""),(2,3,"是→继续"),(2,4,"否→拒绝"),(3,4,"是→登记"),(3,5,"否→拒绝"),(4,5,"")]
)

# 4: 小智AI全链路
make_flowchart(os.path.join(IMG_DIR, 'fc4_xiaozhi.png'), '微信小程序 → 小智AI服务端 → ESP32 设备控制全链路',
    nodes=[
        ("微信小程序\n用户文字输入: \"中午牛肉面, 晚上包300g饺子\"", False),
        ("POST /api/v1/intake/parse\nLLM+规则引擎解析 → ParseResult (饮食条目+面团需求)", False),
        ("用户确认饮食清单 ?", True),
        ("POST /api/v1/recommendations/coarse-grain\n计算杂粮配比 → Recommendation", False),
        ("用户点击\"开始和面\"\nWebSocket MCP指令 → ESP32", False),
        ("ctl_try_acquire(CTL_XIAOZHI)\n互斥锁获取成功 ?", True),
        ("weight_work(weight) → fwork()\n四阶段自动执行 → 返回完成", False),
    ],
    edges=[(0,1,""),(1,2,""),(2,3,"是"),(3,4,""),(4,5,""),(5,6,"是"),(5,3,"否→\"设备正忙\"")]
)

# 5: 推荐算法
make_flowchart(os.path.join(IMG_DIR, 'fc5_recommender.png'), 'recommend_coarse_grain() 营养推荐算法',
    nodes=[
        ("输入: items(当日食物清单) + dough_request(面团需求)", False),
        ("refined = Σrefined_staple (精制主食摄入量)\ncoarse = Σcoarse_grain (杂粮摄入量)", False),
        ("ratio = 0.20 (基础杂粮比例 20%)", False),
        ("refined >= 300g ?\n→ ratio += 0.15", True),
        ("refined >= 150g ?\n→ ratio += 0.10", True),
        ("coarse >= 100g ?\n→ ratio -= 0.05", True),
        ("digestive_risk ?\n→ ratio = min(ratio, 0.20)", True),
        ("ratio = clamp(ratio, 0.10, 0.40)\nround(ratio / 0.05) × 0.05", False),
        ("water = total×0.24, flour = total-water\ncoarse_grain = flour × ratio", False),
        ("输出: Recommendation\n{flour_weight, water_weight, coarse_grain_weight, ratio, reason}", False),
    ],
    edges=[(0,1,""),(1,2,""),(2,3,""),(3,4,"是"),(3,5,"否"),(4,5,"是"),(5,6,"是"),(5,7,"否"),(6,7,""),(7,8,""),(8,9,"")]
)

# 6: 流量计
make_flowchart(os.path.join(IMG_DIR, 'fc6_flowmeter.png'), '水泵加水——YF-S201流量计闭环脉冲计数控制',
    nodes=[
        ("Task2 水泵加水\n输入: water_weight_g (配方计算水量)", False),
        ("target = ceil(water_weight_g × 0.45)\npulse_count = 0 (清零计数器)", False),
        ("ledc_pwm_set_state(1)\n开水泵, GPIO12输出50%占空比PWM", False),
        ("pulse_count < target ?\n轮询等待脉冲达标", True),
        ("[GPIO37上升沿中断 ISR]\npulse_count++ (IRAM中的轻量级ISR)", False),
        ("pulse_count >= target\nledc_pwm_set_state(0) 关断水泵", False),
    ],
    edges=[(0,1,""),(1,2,""),(2,3,""),(3,4,"是→vTaskDelay(10ms)"),(3,5,"否→达标"),(4,3,"")]
)

# 7: 触控屏
make_flowchart(os.path.join(IMG_DIR, 'fc7_touch.png'), '触控屏启动和面流程 action_star_mixer()',
    nodes=[
        ("action_star_mixer()\n用户点击触控屏\"启动\"按钮", False),
        ("w = get_var_dough_weight()\n读取UI当前面团重量值", False),
        ("ctl_try_acquire(CTL_TOUCH)\n互斥锁获取成功 ?", True),
        ("set_var_motor_running(true)\n更新UI状态为\"运行中\"", False),
        ("xTaskCreate(motor_task)\n创建独立FreeRTOS任务 (堆栈8KB, 优先级5)", False),
        ("motor_task → weight_work(w)\n→ fwork() 四阶段自动执行", False),
    ],
    edges=[(0,1,""),(1,2,""),(2,3,"是"),(2,3,"否→拒绝,不响应"),(3,4,""),(4,5,"")]
)

# ═══════════════════ Docx 生成 ═══════════════════
doc = Document()
S = doc.styles['Normal']
S.font.name = '宋体'; S.font.size = Pt(11)
S.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs: r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def P(text):
    p = doc.add_paragraph(text); p.paragraph_format.line_spacing = 1.5

def B(text):
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_before = Pt(10)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11)

def T(hds, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(hds))
    table.style = 'Light Grid Accent 1'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(hds):
        c = table.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph()

def IO(rows):
    T(['变量名','类型','分类','含义'], rows)

def IMG(name, width=6.3):
    path = os.path.join(IMG_DIR, name)
    if os.path.exists(path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(); r.add_picture(path, width=Inches(width))
    doc.add_paragraph()

# ═══════════════════ 完整文档内容 ═══════════════════

H('2.1 整体介绍', level=1)
P('本系统以 ESP32-S3 为主控芯片，构建智能全自动和面机。系统采用四层架构：用户交互层、智能服务层、主控决策层与硬件执行层。')
P('触控屏（LVGL/EEZ Studio）、网页控制面板（HTTP REST API）和微信小程序三种用户入口通过统一的互斥锁机制共享硬件资源。微信小程序端用户以文字或语音描述当日饮食情况和面食需求，经小智 AI 服务端（xiaozhi-server）调用 DeepSeek 大语言模型和营养推荐引擎解析后，生成面团配方的结构化参数，通过 WebSocket 长连接下发至 ESP32 设备端。设备端运行 FreeRTOS 多任务系统，核心工作流引擎根据接收到的参数，按配方比例驱动 DRV8870 水泵、无刷电调、双舵机、步进电机和霍尔流量计等执行机构，完成面粉加入、加水、研磨、加料、搅拌的全自动流程。')

H('系统整体框图', level=2)
T(['层次','组成模块','关键技术'],
  [['用户交互层','触控屏（LVGL+EEZ Studio）\n网页控制面板（HTTP REST）\n微信小程序（小智 AI 入口）','三端互斥锁（ctl_mutex）\nFreeRTOS 多任务调度'],
   ['智能服务层','小智 AI 服务端（xiaozhi-server）\nASR / LLM / TTS 引擎\n营养推荐引擎（nutrition-service）','DeepSeek 大语言模型\nWebSocket 实时通信\nMCP 工具调用协议'],
   ['主控决策层','ESP32-S3（240MHz Xtensa LX7）\nFreeRTOS 实时操作系统\n核心工作流引擎（function.c）','LEDC 8通道全满调度\nI2S 全双工音频管线\nOpus + ESP-SR 唤醒词引擎'],
   ['硬件执行层','DRV8870 水泵（GPIO12/13 PWM）\n无刷电调 ×2（GPIO15/9 ESC）\n双舵机 + 步进电机\nYF-S201 流量计（GPIO37）','5kHz PWM 电机调速\n50Hz ESC 油门控制\n闭环脉冲计数水量控制']])
P('各层之间通过明确的接口通信：用户交互层经 ctl_mutex 统一调度后调用核心工作流引擎；智能服务层通过 WebSocket/MCP 下发结构化指令至设备端；主控决策层通过 LEDC/I2S/GPIO 驱动硬件执行层。层次间低耦合、层内高内聚的设计便于后续独立升级任一模块。')

H('2.3 软件系统介绍', level=1)
H('2.3.1 软件总体设计', level=2)
P('ESP32-S3 搭载 FreeRTOS 实时操作系统，通过抢占式多任务调度管理图形界面渲染、WiFi 网络通信、HTTP/TCP 指令处理、音频语音交互及电机控制工作流等并行任务。任务间通过互斥信号量、消息队列和事件组进行同步，确保多任务并发下的实时响应能力。')
P('ESP32-S3 内部资源在设计中得到充分利用：4 个 LEDC 定时器与全部 8 个通道被分配至 7 路独立 PWM 输出；I2S0 外设配置为全双工模式，以 4 个 GPIO 同时驱动 INMP441 麦克风输入和 MAX98357A 扬声器输出。在 240MHz 单核 Xtensa LX7 处理器上，通过 FreeRTOS 时间片轮转调度，将 CPU 密集型任务与实时性任务分配到不同优先级的任务中，保证关键控制路径的确定性和语音交互的低延迟。')

# ═══════════════════ 2.3.2 ═══════════════════
H('2.3.2 软件各模块介绍', level=1)

# 模块一
H('一、核心工作流引擎（function.c）', level=2)
P('工作流引擎接收面团重量参数，按预设配方比例计算各物料目标值，分四阶段顺序执行全自动和面流程。')
IMG('fc1_stack.png', 6.3)
IMG('fc2_weightwork.png', 6.3)
B('关键输入变量：')
IO([['weight','uint32_t','输入','面团总重量（g），来自触控屏/HTTP/微信小程序'],
    ['k_flour~k_salt','float','常量(5组)','配方比例常数：120/267.5, 65/267.5, 80/267.5, 2/267.5, 0.5/267.5']])
B('关键输出变量：')
IO([['flour','int','输出','面粉电机运行时长（ms）'],
    ['water','int','输出','水泵运行时长/目标脉冲数'],
    ['grain','int','输出','研磨运行总时长（ms）'],
    ['yeast','int','输出','酵母加料循环次数'],
    ['salt','int','输出','盐加料循环次数']])

# 模块二
H('二、三端互斥控制（ctl_mutex.c）', level=2)
P('系统支持触控屏、HTTP 网页、小智 AI 三方控制。本模块基于 FreeRTOS 互斥信号量实现非阻塞互斥调度，确保同一时刻仅有一端操作硬件。采用非阻塞策略（xSemaphoreTake 超时为0）而非阻塞等待，避免触控屏 UI 线程或 HTTP 请求被长时间挂起导致用户体验冻结。')
IMG('fc3_mutex.png', 6.3)
B('控制源接入点：')
T(['控制源','枚举值','获取锁位置','释放锁位置'],
  [['触控屏','CTL_TOUCH','action_star_mixer()','action_stop()'],
   ['HTTP 网页','CTL_HTTP','api_start_post_handler()','api_stop_post_handler()'],
   ['小智 AI','CTL_XIAOZHI','handle_mcp_command()','fstop()']])
B('关键变量：')
IO([['s_mutex','SemaphoreHandle_t','内部','FreeRTOS 互斥信号量'],
    ['s_owner','ctl_source_t','内部','NONE/TOUCH/HTTP/XIAOZHI'],
    ['src (入参)','ctl_source_t','输入','请求控制权的来源标识'],
    ['返回值','bool','输出','true=获取成功, false=已被占用']])

# 模块三
H('三、电机与执行机构驱动', level=2)
P('本系统共驱动 4 类、7 路独立执行机构，全部通过 ESP32-S3 LEDC 硬件外设产生 PWM 信号，经 4 个定时器与 8 个独立通道统一管理，资源利用率 100%。')
B('LEDC 通道分配表：')
T(['Timer','频率','分辨率','Ch','GPIO','执行机构','信号类型'],
  [['0','5kHz','10-bit','CH3','12','水泵 DRV8870 IN1','PWM 可变占空比'],
   ['0','5kHz','10-bit','CH4','13','水泵 DRV8870 IN2','固定 LOW'],
   ['1','1kHz','10-bit','CH1','4','步进电机 STP','脉冲序列'],
   ['2','50Hz','14-bit','CH0','8','旋转舵机','角度-脉宽映射'],
   ['2','50Hz','14-bit','CH2','3','击打舵机','角度-脉宽映射'],
   ['3','50Hz','14-bit','CH6','15','面粉 ESC','油门-脉宽映射'],
   ['3','50Hz','14-bit','CH7','9','搅拌 ESC','油门-脉宽映射']])
P('技术难点：选择 10-bit 分辨率（1024 占空比等级）而非更高位宽的原因是 LEDC 频率可达性约束。ESP32-S3 的 LEDC 时钟源为 80MHz APB 时钟。若采用 13-bit（8192 等级），理论最高频率仅为 9.8kHz，无法达到目标 5kHz。10-bit 模式下最高频率为 78kHz，5kHz 通过分数分频器精确实现。')
B('关键变量：')
IO([['state (入参)','int','输入','0=关闭, 1=开启50%占空比'],
    ['duty','uint32_t','输出','目标占空比值（0 或 511）'],
    ['PWM_FREQ_HZ','宏','常量','50000 (5kHz, DRV8870最佳频段)']])

# 模块四
H('四、小智 AI 智能交互', level=2)
P('用户通过微信小程序聊天界面描述当日饮食和面食需求，经小智 AI 服务端解析后，通过 MCP 协议下发结构化控制指令至 ESP32。服务端由 FastAPI 营养推荐引擎（nutrition-service）和 WebSocket 网关（xiaozhi-server）两部分组成。')
IMG('fc4_xiaozhi.png', 6.3)
IMG('fc5_recommender.png', 6.3)
B('MCP 指令映射：')
T(['用户意图','MCP tool','设备端函数','参数'],
  [['开始和面','start_mixer','weight_work(weight)','weight: 面团总重量(g)'],
   ['紧急停止','stop_mixer','fstop()','无'],
   ['推出面团','push_out','push_and_out(1)','direction: 1'],
   ['退回','push_back','push_and_out(0)','direction: 0']])
B('关键输入变量：')
IO([['text','string','输入','用户当日饮食+面食需求自然语言'],
    ['confirmed_items','list[FoodItem]','输入','用户确认后的饮食条目列表']])
B('关键输出变量：')
IO([['Recommendation.ratio','float','输出','杂粮比例（10%~40%,步进5%）'],
    ['Recommendation.flour','int','输出','普通面粉克重（g）'],
    ['Recommendation.coarse_grain','int','输出','杂粮粉克重（g）'],
    ['Recommendation.water','int','输出','水量（g/mL）'],
    ['MCP JSON','JSON','输出','下发给 ESP32 的结构化控制指令']])

# 模块五
H('五、流量计闭环水量控制', level=2)
P('将传统定时开环加水（精度 ±20%）升级为 YF-S201 霍尔流量传感器闭环脉冲计数控制（精度 ±2~5%）。传感器输出 450 脉冲/L，GPIO37 配置为上升沿中断输入，IRAM 中的轻量级 ISR 仅执行 pulse_count++。未采用 PID 补偿——水泵关断后管路残余水量在传感器固有误差范围内。')
IMG('fc6_flowmeter.png', 6.3)
B('关键变量：')
IO([['pulse_count','volatile uint32_t','全局','IRAM ISR 中的脉冲累加计数器'],
    ['target','uint32_t','局部','目标脉冲数 = ceil(水量mL × 0.45)'],
    ['water_weight_g','int','输入','配方计算出的水量（g，即 mL）']])

# 模块六
H('六、LVGL 触控屏与 HTTP 服务器', level=2)
P('触控屏采用 LVGL v9 + EEZ Studio，EEZ 生成页面布局，action.cpp 处理按钮事件。HTTP 服务器（端口 80）内嵌网页固件，提供 REST API。两个模块在执行电机控制前均需获取互斥锁。')
IMG('fc7_touch.png', 6.3)
B('关键变量：')
IO([['dough_weight','int32_t','全局','UI 当前面团重量值（50~1000g）'],
    ['motor_running','bool','全局','UI 电机运行状态指示'],
    ['g_motor_task_handle','TaskHandle_t','内部','电机任务 FreeRTOS 句柄']])

# ═══════════ 2.3.3 ═══════════
H('2.3.3 软件技术指标', level=1)
T(['指标','数值','指标','数值'],
  [['操作系统','FreeRTOS（抢占式多任务）','音频编码','Opus（16kHz/单声道/60ms帧）'],
   ['固件框架','ESP-IDF v6.0','唤醒词引擎','ESP-SR MultiNet'],
   ['编程语言','C(驱动层)/C++(应用层+音频)','服务端模型','DeepSeek/ChatGLM'],
   ['图形框架','LVGL v9 + EEZ Studio','HTTP服务','内嵌,端口80,REST JSON'],
   ['LEDC占用','8/8通道,4/4定时器,100%','WiFi模式','AP(802.11n,20MHz)'],
   ['固件体积','2.7MB(分区3MB,14%空闲)','I2S模式','全双工(4 GPIO双设备)'],
   ['GPIO占用','34/49 (69%)','PWM输出','7路独立(LEDC)']])

H('2.3.4 完整引脚与片上资源分配', level=1)
P('（此处插入《ESP32-S3完整资源分配表.docx》中的五个表格。）')

out = os.path.join(OUT_DIR, '参赛文档_软件系统介绍_v5.docx')
doc.save(out)
print(f'\nOK: {out}')
