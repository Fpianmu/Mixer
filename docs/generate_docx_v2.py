#!/usr/bin/env python3
"""生成参赛文档 docx v2 — 含流程图 + I/O变量表"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def P(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5
    for r in p.runs:
        r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def B(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11); r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def FLOW(lines):
    """绘制文本流程图"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.2
    for line in lines:
        r = p.add_run(line + '\n')
        r.font.name = 'Consolas'
        r.font.size = Pt(9)

def T(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph()

def IO(rows):
    """输入/输出变量表"""
    T(['变量名', '类型', '分类', '含义'], rows)

# ═══════════════════════════════════════
H('2.3.2 软件各模块介绍', level=1)

# ── 模块一：核心工作流引擎 ──
H('一、核心工作流引擎（function.c）', level=2)

P('工作流引擎接收面团重量参数，按预设配方比例计算各物料目标值，分四阶段顺序执行全自动和面流程。')

B('函数调用流程：')
FLOW([
    'app_main()                                    ',
    '  └─ init_all()                               ',
    '       ├─ ledc_pwm_set_state(0)   [水泵PWM初始化]',
    '       ├─ servo_init/servo2_init  [双舵机初始化] ',
    '       ├─ stepper_init()          [步进电机初始化]',
    '       ├─ esc_init(&esc1)         [面粉电调初始化]',
    '       ├─ esc_init(&esc2)         [搅拌电调初始化]',
    '       ├─ wifi_init_softap()      [WiFi AP启动]  ',
    '       ├─ http_server_start()     [HTTP服务器]    ',
    '       ├─ xz_init_wrapper         [小智AI语音]    ',
    '       ├─ bsp_lcd_init()          [LCD初始化]     ',
    '       └─ ui_init()               [EEZ UI加载]    ',
    '  └─ while(1) { ui_tick() }       [主循环10ms]    ',
])

B('weight_work(weight) 核心流程：')
FLOW([
    'weight_work(300)                                  ',
    '  │                                                 ',
    '  ├─ flour = (int)((300×k_flour×1000)/2.8324)     ',
    '  ├─ water = (int)((300×k_water×1000)/28.9)       ',
    '  ├─ grain = (int)((300×k_grain×1000×60)/4.43)    ',
    '  ├─ yeast = (int)((300×k_yeast)/0.3)   [循环次数] ',
    '  ├─ salt  = (int)((300×k_salt)/0.3)    [循环次数] ',
    '  │                                                 ',
    '  └─ fwork(flour, water, 500, grain, 400000)       ',
    '       │                                             ',
    '       ├─ Task1 面粉加入                             ',
    '       │   [esc1 40%+磁铁阀ON] → delay → OFF         ',
    '       │                                             ',
    '       ├─ Task2 水泵加水                             ',
    '       │   [DRV8870 ON] → 脉冲计数达标 → OFF         ',
    '       │                                             ',
    '       ├─ Task3 研磨+加料                            ',
    '       │   ├─ 酵母×yeast: servo(145°)+servo2(65°)   ',
    '       │   ├─ 盐×salt:   servo(55°)+servo2(60°)    ',
    '       │   └─ 研磨: 60s ON / 30s OFF 间歇循环       ',
    '       │                                             ',
    '       └─ Task4 搅拌                                 ',
    '           [esc2 40%] → delay(400s) → OFF            ',
])

B('关键输入变量：')
IO([
    ['weight',   'uint32_t', '输入', '面团总重量（g），来自触控屏/HTTP/微信小程序'],
    ['k_flour',  'float',    '输入', '面粉配方比例 = 120/267.5'],
    ['k_water',  'float',    '输入', '水配方比例 = 65/267.5'],
    ['k_grain',  'float',    '输入', '杂粮配方比例 = 80/267.5'],
    ['k_yeast',  'float',    '输入', '酵母配方比例 = 2/267.5'],
    ['k_salt',   'float',    '输入', '盐配方比例 = 0.5/267.5'],
])

B('关键输出变量：')
IO([
    ['flour',    'int',      '输出', '面粉电机运行时长（ms），传入 fwork 的 duration1'],
    ['water',    'int',      '输出', '水泵运行时长（ms），传入 fwork 的 duration2'],
    ['grain',    'int',      '输出', '研磨运行总时长（ms），传入 fwork 的 duration4'],
    ['yeast',    'int',      '输出', '酵母加料循环次数'],
    ['salt',     'int',      '输出', '盐加料循环次数'],
    ['esc1/esc2','esc_handle_t*','全局','电调句柄，保存校准状态和油门值'],
])

# ── 模块二：互斥控制 ──
H('二、三端互斥控制（ctl_mutex.c）', level=2)

P('系统支持触控屏、HTTP 网页、小智 AI 三方控制。三方共享同一组硬件执行机构，本模块基于 FreeRTOS 互斥信号量实现非阻塞互斥调度。')

B('ctl_try_acquire(src) 流程图：')
FLOW([
    'ctl_try_acquire(CTL_TOUCH)                         ',
    '  │                                                  ',
    '  ├─ xSemaphoreTake(mutex, 0)   [非阻塞尝试]        ',
    '  │   └─ 失败 → 返回 false     [已被占用]            ',
    '  │                                                  ',
    '  ├─ s_owner != NONE && s_owner != src              ',
    '  │   └─ 释放锁 → 返回 false   [其他源占用中]        ',
    '  │                                                  ',
    '  └─ s_owner = CTL_TOUCH → 释放锁 → 返回 true       ',
])

B('各控制源接入点：')
T(['控制源', '枚举值', '获取锁位置', '释放锁位置'],
  [['触控屏', 'CTL_TOUCH', 'action_star_mixer()', 'action_stop()'],
   ['HTTP 网页', 'CTL_HTTP', 'api_start_post_handler()', 'api_stop_post_handler()'],
   ['小智 AI', 'CTL_XIAOZHI', 'handle_mcp_command()', 'fstop()']])

B('关键变量：')
IO([
    ['s_mutex',   'SemaphoreHandle_t', '内部', 'FreeRTOS 互斥信号量，保护 s_owner 读写'],
    ['s_owner',   'ctl_source_t',      '内部', '当前控制权归属：NONE/TOUCH/HTTP/XIAOZHI'],
    ['src (入参)', 'ctl_source_t',      '输入', '请求控制权的来源标识'],
    ['返回值',     'bool',              '输出', 'true=获取成功，false=已被占用'],
])

# ── 模块三：电机驱动 ──
H('三、电机与执行机构驱动（ledc_pwm / esc_controller / servo / stepper）', level=2)

P('本系统共驱动 4 类、7 路独立执行机构，全部通过 ESP32-S3 LEDC 硬件外设产生 PWM 信号，经 4 个定时器与 8 个独立通道统一管理。')

B('LEDC 通道分配：')
T(['Timer', '频率', '分辨率', 'Ch', 'GPIO', '执行机构', '信号类型'],
  [['0','5kHz','10-bit','CH3','12','水泵 DRV8870 IN1','PWM 可变占空比'],
   ['0','5kHz','10-bit','CH4','13','水泵 DRV8870 IN2','固定 LOW'],
   ['1','1kHz','10-bit','CH1','4','步进电机 STP','脉冲序列'],
   ['2','50Hz','14-bit','CH0','8','旋转舵机','角度-脉宽映射'],
   ['2','50Hz','14-bit','CH2','3','击打舵机','角度-脉宽映射'],
   ['3','50Hz','14-bit','CH6','15','面粉 ESC','油门-脉宽映射'],
   ['3','50Hz','14-bit','CH7','9','搅拌 ESC','油门-脉宽映射']])

B('ledc_pwm_set_state(state) 流程图：')
FLOW([
    'ledc_pwm_set_state(1)                               ',
    '  │                                                  ',
    '  ├─ [首次调用] 一次性初始化                          ',
    '  │   ├─ ledc_timer_config(Timer0, 5kHz, 10bit)     ',
    '  │   ├─ ledc_channel_config(CH3 → GPIO12)           ',
    '  │   └─ ledc_channel_config(CH4 → GPIO13)           ',
    '  │                                                  ',
    '  ├─ state==0 → duty=0        [关闭: 两路占空比 0%]  ',
    '  └─ state==1 → duty=511      [开启: CH3=50%, CH4=0%]',
    '       │                                             ',
    '       ├─ ledc_set_duty(LOW_SPEED, CH3, duty)        ',
    '       └─ ledc_update_duty(LOW_SPEED, CH3)           ',
])

B('关键变量：')
IO([
    ['state (入参)', 'int', '输入', '0=关闭, 1=开启 50%占空比'],
    ['duty_ch0',     'uint32_t','输出','CH3 目标占空比值（0 或 511）'],
    ['PWM_FREQ_HZ',  '#define', '常量','50000（5kHz，DRV8870 最佳频段）'],
    ['DUTY_50_PERCENT','#define','常量','511（1023 总等级的一半）'],
])

# ── 模块四：XZ AI ──
H('四、小智 AI 智能交互（xiaozhi / nutrition-service）', level=2)

P('用户通过微信小程序聊天界面描述当日饮食和面食需求，经小智 AI 服务端解析后，通过 MCP 协议下发结构化控制指令至 ESP32。')

B('数据流与函数调用链：')
FLOW([
    '微信小程序 (用户文字输入)                               ',
    '  │                                                    ',
    '  ├─ POST /api/v1/intake/parse                        ',
    '  │   → NutritionService.parse_intake()                ',
    '  │   → LLM + 规则引擎 → ParseResult {items, dough_request}',
    '  │                                                    ',
    '  ├─ 用户确认/修改饮食清单                              ',
    '  │   POST /api/v1/recommendations/coarse-grain       ',
    '  │   → NutritionService.confirm_and_recommend()       ',
    '  │   → recommender.recommend_coarse_grain()          ',
    '  │   → Recommendation {flour, water, coarse_grain}    ',
    '  │                                                    ',
    '  └─ 用户点击"开始和面"                                 ',
    '      → WebSocket MCP: {"tool":"start_mixer","args":{}}',
    '      → ESP32: handle_mcp_command()                    ',
    '      → ctl_try_acquire(CTL_XIAOZHI)                  ',
    '      → weight_work(weight)                            ',
    '      → fwork(...) [四阶段自动执行]                     ',
])

B('MCP 指令映射：')
T(['用户意图', 'MCP tool', '设备端函数', '参数'],
  [['开始和面', 'start_mixer', 'weight_work(weight)', 'weight: 面团总重量(g)'],
   ['紧急停止', 'stop_mixer', 'fstop()', '无'],
   ['推出面团', 'push_out', 'push_and_out(1)', 'direction: 1 (CCW)'],
   ['退回', 'push_back', 'push_and_out(0)', 'direction: 0 (CW)']])

B('推荐算法核心逻辑：')
FLOW([
    'recommend_coarse_grain(items, dough_request)       ',
    '  │                                                  ',
    '  ├─ refined = sum(items.refined_staple)  [今日精制主食]',
    '  ├─ coarse  = sum(items.coarse_grain)    [今日杂粮]    ',
    '  │                                                  ',
    '  ├─ ratio = 0.20                    [基础 20%]       ',
    '  │   refined≥300 → ratio += 0.15                     ',
    '  │   refined≥150 → ratio += 0.10                     ',
    '  │   coarse≥100  → ratio -= 0.05                     ',
    '  │   digestive_risk → ratio = min(ratio, 0.20)       ',
    '  │   ratio = clamp(ratio, 0.10, 0.40)                ',
    '  │                                                  ',
    '  ├─ water = total × 0.24              [水量 = 总重×24%]',
    '  ├─ flour = total - water             [面粉总重]       ',
    '  └─ coarse_grain = flour × ratio      [杂粮粉重]       ',
])

B('关键输入变量：')
IO([
    ['text',        'string',   '输入', '用户当日饮食+面食需求自然语言描述'],
    ['user_id',     'string',   '输入', '用户标识（由微信 openid 或预设值）'],
    ['confirmed_items','list[FoodItem]','输入','用户确认后的饮食条目列表'],
    ['dough_request','DoughRequest','输入','面团需求（目标食物类型、总重量）'],
])

B('关键输出变量：')
IO([
    ['ParseResult.items',   'list[FoodItem]','输出','解析出的食物条目（meal, weight, category）'],
    ['Recommendation.ratio', 'float',  '输出', '杂粮比例（10%~40%，步进 5%）'],
    ['Recommendation.flour', 'int',    '输出', '普通面粉克重（g）'],
    ['Recommendation.coarse_grain','int','输出','杂粮粉克重（g）'],
    ['Recommendation.water', 'int',    '输出', '水量（g/mL）'],
    ['MCP tool+args',       'JSON',   '输出', '下发给 ESP32 的结构化控制指令'],
])

# ── 模块五：流量计 ──
H('五、流量计闭环水量控制（function.c + GPIO 中断）', level=2)

P('将传统定时开环加水（精度 ±20%）升级为 YF-S201 霍尔流量传感器闭环脉冲计数控制（精度 ±2~5%）。')

B('脉冲计数控制流程：')
FLOW([
    'Task2: 水泵加水 (闭环)                              ',
    '  │                                                  ',
    '  ├─ target = (water_weight_mL × 450) / 1000       ',
    '  │   [目标脉冲数 = 水量mL × 0.45]                   ',
    '  │                                                  ',
    '  ├─ pulse_count = 0                    [清零]       ',
    '  ├─ ledc_pwm_set_state(1)              [开水泵]      ',
    '  │                                                  ',
    '  ├─ while (pulse_count < target):      [轮询等待]    ',
    '  │     vTaskDelay(10ms)                              ',
    '  │                                                  ',
    '  └─ pulse_count ≥ target → ledc_pwm_set_state(0)   ',
    '                                  [关断水泵]          ',
    '',
    'GPIO37 ISR (IRAM):                                  ',
    '  └─ pulse_count++                   [仅累加计数]    ',
])

B('关键变量：')
IO([
    ['pulse_count', 'volatile uint32_t', '全局', 'IRAM ISR 中的脉冲累加计数器'],
    ['target',      'uint32_t',    '局部', '目标脉冲数 = ceil(水量mL × 0.45)'],
    ['water_weight_g','int',       '输入', '配方计算出的水量（g，即 mL）'],
    ['flow_rate',   'float',       '输出', '瞬时流量 = (freq_Hz / 7.5) L/min'],
    ['total_volume','float',       '输出', '累计水量 = pulse_count / 450 L'],
])

# ── 模块六：辅助 ──
H('六、辅助模块（LVGL 触控屏 / HTTP 服务器）', level=2)

P('触控屏采用 LVGL v9 + EEZ Studio，EEZ 生成页面布局（screens.c/ui.c），action.cpp 处理按钮事件并映射至工作流函数。HTTP 服务器（端口 80）内嵌静态网页固件，提供 REST API：POST /api/start（JSON {"weight":N}）、POST /api/stop、GET /api/status。两个模块在执行电机控制前均需通过 ctl_try_acquire 获取互斥锁。')

B('action_star_mixer 流程：')
FLOW([
    'action_star_mixer(lv_event_t *e)                      ',
    '  │                                                    ',
    '  ├─ w = get_var_dough_weight()     [读取UI重量值]      ',
    '  ├─ ctl_try_acquire(CTL_TOUCH)     [尝试获取互斥锁]    ',
    '  │   └─ 失败 → return              [系统忙，拒绝操作]   ',
    '  ├─ set_var_motor_running(true)    [更新UI状态]        ',
    '  └─ xTaskCreate(motor_task, ...)   [创建独立任务执行]   ',
    '       └─ weight_work(w) → fwork()  [调用工作流引擎]     ',
])

B('关键变量：')
IO([
    ['dough_weight',   'int32_t',    '全局', 'UI 设置的当前面团重量值（50~1000g）'],
    ['motor_running',  'bool',       '全局', 'UI 中电机的运行状态指示'],
    ['g_motor_task_handle','TaskHandle_t','内部','电机任务的 FreeRTOS 句柄'],
    ['g_current_weight','uint32_t',  '内部', 'HTTP 模块存储的当前面团重量'],
])

# ═══════════════════════════════════
H('2.3.3 软件技术指标', level=1)

T(['指标', '数值', '指标', '数值'],
  [['操作系统', 'FreeRTOS（抢占式多任务）','音频编码','Opus（16kHz/单声道/60ms帧）'],
   ['固件框架', 'ESP-IDF v6.0 (CMake+Ninja)','唤醒词引擎','ESP-SR MultiNet'],
   ['编程语言', 'C（驱动层）/ C++（应用层+音频）','服务端模型','DeepSeek / ChatGLM'],
   ['图形框架', 'LVGL v9 + EEZ Studio','HTTP 服务','内嵌服务器, 端口 80, REST JSON'],
   ['LEDC 占用', '8/8 通道, 4/4 定时器, 100%','WiFi 模式','AP 模式 (802.11n, 20MHz)'],
   ['固件体积', '2.7MB（工厂分区 3MB, 14%空闲）','I2S 模式','全双工（4 GPIO 双设备）'],
   ['GPIO 占用', '34/49 (69%)','PWM 输出','7 路独立（LEDC）'],
])

# ═══════════════════════════════════
H('2.3.4 完整引脚与片上资源分配', level=1)
P('（此处插入《ESP32-S3完整资源分配表.docx》中的五个表格：片上外设占用、LEDC通道详细分配、I2S音频引脚、完整GPIO引脚分配、资源利用率总览。）')

# ── Save ──
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), '参赛文档_软件系统介绍_v2.docx')
doc.save(out)
print(f'OK: {out}')
