#!/usr/bin/env python3
"""生成参赛文档 docx v3 — 表格化流程图"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()
S = doc.styles['Normal']
S.font.name = '宋体'; S.font.size = Pt(11)
S.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def P(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5

def B(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_before = Pt(10)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11)

def T(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph()

def IO(rows):
    T(['变量名', '类型', '分类', '含义'], rows)

# ── Flowchart builder using tables ──
def FC(title, boxes, connections=None):
    """用表格模拟流程图。boxes: [(text1,), (text2,), ...] 每行一个或多个并排box"""
    B(title)

    # 找最大列数
    max_cols = max(len(row) for row in boxes)
    total_rows = len(boxes)

    table = doc.add_table(rows=total_rows, cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表格边框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')

    for ri, row_boxes in enumerate(boxes):
        row = table.rows[ri]
        row.height = Inches(0.45)
        for ci in range(max_cols):
            cell = row.cells[ci]
            # 设置单元格边框和背景
            tcPr = cell._tc.get_or_add_tcPr()

            if ci < len(row_boxes):
                text = row_boxes[ci]
                # 填充背景色
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F0FE" w:val="clear"/>')
                tcPr.append(shading)
            else:
                text = ''
                # 空白单元格透明
                pass

            # 设置边框
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:top w:val="single" w:sz="4" w:color="4472C4"/>'
                f'<w:left w:val="single" w:sz="4" w:color="4472C4"/>'
                f'<w:bottom w:val="single" w:sz="4" w:color="4472C4"/>'
                f'<w:right w:val="single" w:sz="4" w:color="4472C4"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(borders)

            # 设置文字
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            for line in text.split('\n'):
                if p.text:
                    p = cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(line.strip())
                r.font.size = Pt(8)
                r.font.name = 'Consolas'

    # 隐藏表格最外层边框，只保留cell边框
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')

    doc.add_paragraph()


# ═══════════════════════════════════════
H('2.3.2 软件各模块介绍', level=1)

# ── 模块一 ──
H('一、核心工作流引擎（function.c）', level=2)
P('工作流引擎接收面团重量参数，按预设配方比例计算各物料目标值，分四阶段顺序执行全自动和面流程。')

FC('函数调用层次', [
    ['app_main()'],
    ['init_all()'],
    ['ledc_pwm\n水泵PWM', 'servo_init\n旋转舵机', 'servo2_init\n击打舵机', 'stepper_init\n步进电机', 'esc_init\n双路电调'],
    ['wifi_init\n无线网络', 'http_server\n网页控制', 'xz_init\n小智AI语音', 'lcd_init\n显示屏', 'ui_init\n触控界面'],
    ['while(1) { ui_tick() }    主循环 10ms 驱动LVGL'],
])

FC('weight_work(weight) → fwork() 核心工作流', [
    ['weight_work(300)   输入: 面团重量'],
    ['flour=计算配方    water=计算配方    grain=计算配方    yeast/salt=计算循环次数'],
    ['fwork(flour, water, grain, ...)    四阶段顺序执行'],
    ['Task1 面粉加入\nESC1 40%油门\n磁铁阀 GPIO18 ON\n延时 → 关闭'],
    ['Task2 水泵加水\nDRV8870 ON\n流量计脉冲计数\n达标 → 关闭'],
    ['Task3 研磨+加料\n研磨GPIO46 ON\n酵母×yeast次: servo145° servo2 65°\n盐×salt次: servo55° servo2 60°\n研磨间歇 ON/OFF 循环'],
    ['Task4 搅拌\nESC2 40%油门\n延时 400s → 关闭'],
    ['和面完成\n返回IDLE状态'],
])

B('关键输入变量：')
IO([
    ['weight','uint32_t','输入','面团总重量（g），来自触控屏/HTTP/微信小程序'],
    ['k_flour','float','常量','面粉配方比例 = 120/267.5'],
    ['k_water','float','常量','水配方比例 = 65/267.5'],
    ['k_grain','float','常量','杂粮配方比例 = 80/267.5'],
    ['k_yeast','float','常量','酵母配方比例 = 2/267.5'],
    ['k_salt','float','常量','盐配方比例 = 0.5/267.5'],
])

B('关键输出变量：')
IO([
    ['flour','int','输出','面粉电机运行时长（ms）'],
    ['water','int','输出','水泵运行时长/目标脉冲数'],
    ['grain','int','输出','研磨运行总时长（ms）'],
    ['yeast','int','输出','酵母加料循环次数'],
    ['salt','int','输出','盐加料循环次数'],
])

# ── 模块二 ──
H('二、三端互斥控制（ctl_mutex.c）', level=2)
P('系统支持触控屏、HTTP 网页、小智 AI 三方控制。三方共享同一组硬件执行机构，本模块基于 FreeRTOS 互斥信号量实现非阻塞互斥调度。')

FC('ctl_try_acquire(src) 控制权获取流程', [
    ['ctl_try_acquire(CTL_TOUCH)    触控屏请求控制权'],
    ['xSemaphoreTake(mutex, 0)    非阻塞尝试获取信号量'],
    ['获取失败？\n返回 false    系统正忙，拒绝操作'],
    ['s_owner ≠ NONE？\n其他源已占用\n返回 false    被 (HTTP/语音) 占用中'],
    ['s_owner = CTL_TOUCH    登记为当前控制者'],
    ['xSemaphoreGive(mutex) → return true    获取成功，允许执行'],
])

B('控制源接入点：')
T(['控制源','枚举值','获取锁位置','释放锁位置'],
  [['触控屏','CTL_TOUCH','action_star_mixer()','action_stop()'],
   ['HTTP 网页','CTL_HTTP','api_start_post_handler()','api_stop_post_handler()'],
   ['小智 AI','CTL_XIAOZHI','handle_mcp_command()','fstop()']])

B('关键变量：')
IO([
    ['s_mutex','SemaphoreHandle_t','内部','FreeRTOS 互斥信号量'],
    ['s_owner','ctl_source_t','内部','NONE/TOUCH/HTTP/XIAOZHI'],
    ['src','ctl_source_t','输入','请求控制权的来源标识'],
    ['返回值','bool','输出','true=获取成功, false=已被占用'],
])

# ── 模块三 ──
H('三、电机与执行机构驱动', level=2)
P('本系统共驱动 4 类、7 路独立执行机构，全部通过 ESP32-S3 LEDC 硬件外设产生 PWM 信号，经 4 个定时器与 8 个独立通道统一管理。')

B('LEDC 通道分配：')
T(['Timer','频率','分辨率','Ch','GPIO','执行机构','信号类型'],
  [['0','5kHz','10-bit','CH3','12','水泵 DRV8870 IN1','PWM 可变占空比'],
   ['0','5kHz','10-bit','CH4','13','水泵 DRV8870 IN2','固定 LOW'],
   ['1','1kHz','10-bit','CH1','4','步进电机 STP','脉冲序列'],
   ['2','50Hz','14-bit','CH0','8','旋转舵机','角度-脉宽映射'],
   ['2','50Hz','14-bit','CH2','3','击打舵机','角度-脉宽映射'],
   ['3','50Hz','14-bit','CH6','15','面粉 ESC','油门-脉宽映射'],
   ['3','50Hz','14-bit','CH7','9','搅拌 ESC','油门-脉宽映射']])

FC('ledc_pwm_set_state(state) 水泵 PWM 控制', [
    ['ledc_pwm_set_state(1)    开启水泵'],
    ['[首次调用]  一次性初始化\nledc_timer_config(Timer0, 5kHz, 10bit)\nledc_channel_config(CH3→GPIO12, CH4→GPIO13)'],
    ['state == 0 ?\nduty = 0 (关闭)\n或\nduty = 511 (50%占空比)'],
    ['ledc_set_duty(CH3, duty)    设置占空比寄存器'],
    ['ledc_update_duty(CH3)    锁存生效'],
    ['GPIO12 输出 5kHz PWM\n50%占空比 → 电机运转'],
])

B('关键变量：')
IO([
    ['state','int','输入','0=关闭, 1=开启 50%占空比'],
    ['duty','uint32_t','输出','目标占空比值（0 或 511）'],
    ['PWM_FREQ_HZ','宏','常量','50000 (5kHz, DRV8870 最佳频段)'],
    ['DUTY_50_PCT','宏','常量','511 (1023总等级的50%)'],
])

# ── 模块四 ──
H('四、小智 AI 智能交互', level=2)
P('用户通过微信小程序聊天界面描述当日饮食和面食需求，经小智 AI 服务端解析后，通过 MCP 协议下发结构化控制指令至 ESP32。')

FC('微信小程序 → 服务端 → ESP32 全链路', [
    ['微信小程序    用户输入："中午牛肉面，晚上包300g饺子"'],
    ['POST /api/v1/intake/parse\nLLM + 规则引擎 → ParseResult\n(FoodItem[] + DoughRequest)'],
    ['用户确认饮食清单\nPOST /api/v1/recommendations/coarse-grain'],
    ['营养推荐引擎计算\nratio = f(refined_weight, coarse_weight)\n→ Recommendation (flour, water, coarse_grain)'],
    ['用户点击 "开始和面"\nWebSocket MCP 指令 → ESP32'],
    ['ESP32: handle_mcp_command()\n→ ctl_try_acquire(CTL_XIAOZHI)'],
    ['获取锁成功？\n否 → 语音反馈 "设备正忙"'],
    ['weight_work(300) → fwork()\n四阶段自动执行 → 完成'],
])

FC('推荐算法: recommend_coarse_grain()', [
    ['输入: items (当日食物清单) + dough_request (面团需求)'],
    ['refined = Σrefined_staple  (今日精制主食摄入量)\ncoarse = Σcoarse_grain     (今日杂粮摄入量)'],
    ['ratio = 0.20  (基础20%)\nrefined≥300g → +15%\nrefined≥150g → +10%\ncoarse≥100g  → -5%\ndigestive_risk → max 20%'],
    ['ratio = clamp(ratio, 0.10, 0.40)\nratio = round(ratio / 0.05) × 0.05'],
    ['water = total × 0.24\nflour = total - water\ncoarse_grain = flour × ratio'],
    ['输出: Recommendation\n{flour_weight_g, water_weight_g,\ncoarse_grain_weight_g, ratio, reason}'],
])

B('MCP 指令映射：')
T(['用户意图','MCP tool','设备端函数','参数'],
  [['开始和面','start_mixer','weight_work(weight)','weight: 面团总重量(g)'],
   ['紧急停止','stop_mixer','fstop()','无'],
   ['推出面团','push_out','push_and_out(1)','direction: 1'],
   ['退回','push_back','push_and_out(0)','direction: 0']])

B('关键输入变量：')
IO([
    ['text','string','输入','用户当日饮食+面食需求自然语言'],
    ['confirmed_items','list[FoodItem]','输入','用户确认后的饮食条目列表'],
    ['dough_request','DoughRequest','输入','面团需求（食物类型、总重量）'],
])

B('关键输出变量：')
IO([
    ['ParseResult.items','list[FoodItem]','输出','食物条目（meal, weight, category）'],
    ['Recommendation.ratio','float','输出','杂粮比例（10%~40%，步进5%）'],
    ['Recommendation.flour','int','输出','普通面粉克重（g）'],
    ['Recommendation.coarse_grain','int','输出','杂粮粉克重（g）'],
    ['Recommendation.water','int','输出','水量（g/mL）'],
    ['MCP JSON','JSON','输出','下发给 ESP32 的结构化控制指令'],
])

# ── 模块五 ──
H('五、流量计闭环水量控制', level=2)
P('将传统定时开环加水（精度 ±20%）升级为 YF-S201 霍尔流量传感器闭环脉冲计数控制（精度 ±2~5%）。')

FC('Task2: 水泵加水（闭环脉冲计数）', [
    ['输入: water_weight_g (来自配方计算)'],
    ['target = ceil(water_weight_g × 0.45)\n(目标脉冲数)'],
    ['pulse_count = 0    清零计数器'],
    ['ledc_pwm_set_state(1)    开水泵, GPIO12输出50%PWM'],
    ['while (pulse_count < target)\nvTaskDelay(10ms) 轮询等待'],
    ['GPIO37 上升沿中断 ISR:\npulse_count++    (仅累加计数)'],
    ['pulse_count ≥ target    脉冲达标'],
    ['ledc_pwm_set_state(0)    关断水泵'],
    ['精度: 300g面团±5%, 500g以上±4%'],
])

B('关键变量：')
IO([
    ['pulse_count','volatile uint32_t','全局','IRAM ISR 中的脉冲累加计数器'],
    ['target','uint32_t','局部','目标脉冲数 = ceil(水量mL × 0.45)'],
    ['water_weight_g','int','输入','配方计算出的水量（g, 即 mL）'],
    ['flow_rate','float','输出','瞬时流量 = (freq_Hz / 7.5) L/min'],
    ['total_volume','float','输出','累计水量 = pulse_count / 450 L'],
])

# ── 模块六 ──
H('六、LVGL 触控屏与 HTTP 服务器', level=2)
P('触控屏采用 LVGL v9 + EEZ Studio，EEZ 生成页面布局，action.cpp 处理按钮事件。HTTP 服务器（端口 80）内嵌网页固件，提供 REST API。')

FC('触控屏启动流程: action_star_mixer()', [
    ['action_star_mixer()    用户点击 "启动" 按钮'],
    ['w = get_var_dough_weight()    读取UI当前重量值'],
    ['ctl_try_acquire(CTL_TOUCH)    尝试获取互斥锁'],
    ['获取成功？\n否 → 返回 (不响应)'],
    ['set_var_motor_running(true)    更新UI状态'],
    ['xTaskCreate(motor_task)    创建独立FreeRTOS任务'],
    ['motor_task: weight_work(w)\n→ fwork() 四阶段执行'],
])

B('关键变量：')
IO([
    ['dough_weight','int32_t','全局','UI 当前面团重量值（50~1000g）'],
    ['motor_running','bool','全局','UI 电机运行状态指示'],
    ['g_motor_task_handle','TaskHandle_t','内部','电机任务 FreeRTOS 句柄'],
])

# ═══════════════════════════════════════
H('2.3.3 软件技术指标', level=1)

T(['指标','数值','指标','数值'],
  [['操作系统','FreeRTOS（抢占式多任务）','音频编码','Opus（16kHz/单声道/60ms帧）'],
   ['固件框架','ESP-IDF v6.0','唤醒词引擎','ESP-SR MultiNet'],
   ['编程语言','C(驱动层) / C++(应用层+音频)','服务端模型','DeepSeek / ChatGLM'],
   ['图形框架','LVGL v9 + EEZ Studio','HTTP服务','内嵌服务器, 端口80, REST JSON'],
   ['LEDC占用','8/8通道, 4/4定时器, 100%','WiFi模式','AP (802.11n, 20MHz)'],
   ['固件体积','2.7MB (工厂分区3MB, 14%空闲)','I2S模式','全双工 (4 GPIO双设备)'],
   ['GPIO占用','34/49 (69%)','PWM输出','7路独立 (LEDC)'],
])

# ── Save ──
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), '参赛文档_软件系统介绍_v3.docx')
doc.save(out)
print(f'OK: {out}')
