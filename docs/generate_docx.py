#!/usr/bin/env python3
"""生成参赛文档 docx — 2.3 软件系统介绍 + 2.1 整体介绍"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 样式设置 ──
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── 辅助函数 ──
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r = p.add_run(text)
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table

# ═══════════════════════════════════════════════════════════════
# 2.1 整体介绍
# ═══════════════════════════════════════════════════════════════
add_heading('2.1 整体介绍', level=1)

add_para('本系统以 ESP32-S3 为主控芯片，构建智能全自动和面机。系统采用四层架构：用户交互层、智能服务层、主控决策层与硬件执行层。', bold_prefix='系统概述：')

add_para('触控屏（LVGL/EEZ Studio）、网页控制面板（HTTP REST API）和微信小程序三种用户入口通过统一的互斥锁机制共享硬件资源。微信小程序端用户以文字或语音描述当日饮食情况和面食需求，经小智 AI 服务端（xiaozhi-server）调用 DeepSeek 大语言模型和营养推荐引擎解析后，生成面团配方的结构化参数，通过 WebSocket 长连接下发至 ESP32 设备端。设备端运行 FreeRTOS 多任务系统，核心工作流引擎根据接收到的参数，按配方比例驱动 DRV8870 水泵、无刷电调、双舵机、步进电机和霍尔流量计等执行机构，完成面粉加入、加水、研磨、加料、搅拌的全自动流程。', bold_prefix='工作流程：')

add_heading('系统整体框图', level=2)

# 框图简化为表格形式
add_table(
    ['层次', '组成模块', '关键技术'],
    [
        ['用户交互层', '触控屏（LVGL+EEZ Studio）\n网页控制面板（HTTP REST）\n微信小程序（小智 AI 入口）', '三端互斥锁（ctl_mutex）\nFreeRTOS 多任务调度'],
        ['智能服务层', '小智 AI 服务端（xiaozhi-server）\nASR / LLM / TTS 引擎\n营养推荐引擎（nutrition-service）', 'DeepSeek 大语言模型\nWebSocket 实时通信\nMCP 工具调用协议'],
        ['主控决策层', 'ESP32-S3（240MHz Xtensa LX7）\nFreeRTOS 实时操作系统\n核心工作流引擎（function.c）', 'LEDC 8通道全满调度\nI2S 全双工音频管线\nOpus + ESP-SR 唤醒词引擎'],
        ['硬件执行层', 'DRV8870 水泵（GPIO12/13 PWM）\n无刷电调 ×2（GPIO15/9 ESC）\n双舵机 + 步进电机\nYF-S201 流量计（GPIO37）', '5kHz PWM 电机调速\n50Hz ESC 油门控制\n闭环脉冲计数水量控制'],
    ]
)

add_para('各层之间通过明确的接口通信：用户交互层经 ctl_mutex 统一调度后调用核心工作流引擎；智能服务层通过 WebSocket/MCP 下发结构化指令至设备端；主控决策层通过 LEDC/I2S/GPIO 驱动硬件执行层。层次间低耦合、层内高内聚的设计便于后续独立升级任一模块。')

# ═══════════════════════════════════════════════════════════════
# 2.3 软件系统介绍
# ═══════════════════════════════════════════════════════════════
add_heading('2.3 软件系统介绍', level=1)

add_heading('2.3.1 软件总体设计', level=2)

add_para('ESP32-S3 搭载 FreeRTOS 实时操作系统，通过抢占式多任务调度管理图形界面渲染、WiFi 网络通信、HTTP/TCP 指令处理、音频语音交互及电机控制工作流等并行任务。任务间通过互斥信号量、消息队列和事件组进行同步，确保多任务并发下的实时响应能力。', bold_prefix='系统架构：')

add_para('ESP32-S3 内部资源在设计中得到充分利用：4 个 LEDC 定时器与全部 8 个通道被分配至 7 路独立 PWM 输出（水泵 DRV8870、步进电机、双舵机、双路 ESC）；I2S0 外设配置为全双工模式，以 4 个 GPIO 同时驱动 INMP441 麦克风输入和 MAX98357A 扬声器输出；2 个硬件 UART、1 路 SPI（LCD）、1 路 I2C（触摸）均被占用。在 240MHz 单核 Xtensa LX7 处理器上，通过 FreeRTOS 时间片轮转调度，将 CPU 密集型任务（Opus 音频编码、LVGL 图形渲染）与实时性任务（电机 PWM 更新、流量计中断计数）分配到不同优先级的任务中，保证关键控制路径的确定性和语音交互的低延迟。', bold_prefix='资源利用：')

add_heading('2.3.2 核心工作流引擎', level=2)

add_para('工作流引擎 function.c 接收面团重量参数，按预设配方比例（k_flour、k_water、k_grain、k_yeast、k_salt）计算各物料目标值，分四阶段顺序执行全自动和面流程。', bold_prefix='功能概述：')

add_table(
    ['阶段', '执行机构', '控制方式', '精度'],
    [
        ['Task1 面粉加入', '面粉 ESC（GPIO15, 50Hz PWM）\n面粉磁铁阀（GPIO18）', '40% 油门, 分段间歇加入', '运行时长由配方计算'],
        ['Task2 水泵加水', 'DRV8870（GPIO12/13, 5kHz PWM）\nYF-S201 流量计（GPIO37 中断）', '50% 占空比 PWM\n流量计脉冲计数关断', '±5%（300g面团）\n±4%（500g以上）'],
        ['Task3 研磨+加料', '研磨继电器（GPIO46）\n旋转舵机（GPIO8, 50Hz）\n击打舵机（GPIO3, 50Hz）', '60s ON/30s OFF 间歇研磨\n舵机定位→击打→归位循环', '循环次数由配方计算'],
        ['Task4 搅拌', '搅拌 ESC（GPIO9, 50Hz）', '40% 油门持续运转', '运行时长由配方计算'],
    ]
)

add_para('配方计算中，五组比例常数均为 float 浮点类型。为确保计算精度，所有中间运算保持浮点，仅在最终赋值时进行 (int) 强制类型转换。以酵母和盐为例，改前 "(int)(weight × k_yeast) / 0.3" 会在除法前截断小数，导致 300g 面团时盐的循环次数被误算为 0；改后 "(int)((weight × k_yeast) / 0.3)" 将截断置于除法之后，盐循环次数修正为 1。该修正消除了早期版本中因 (int) 位置不当导致的加料遗漏。', bold_prefix='精度处理：')

add_para('四阶段采用严格顺序执行而非状态机并发设计。原因在于和面工艺的不可逆性——后续阶段必须在前一阶段完成后才能开始，交替执行将导致加料顺序错误。此设计牺牲了一定的并行性，但保证了食品加工的可重复性。', bold_prefix='设计约束：')

add_heading('2.3.3 三端互斥控制', level=2)

add_para('系统支持触控屏、HTTP 网页、小智 AI 三方控制。三方共享同一组硬件执行机构，若同时创建电机任务将导致硬件冲突。本模块基于 FreeRTOS 互斥信号量实现非阻塞互斥调度，提供 ctl_try_acquire(src)、ctl_release(src)、ctl_get_owner() 三个接口。', bold_prefix='问题描述：')

add_table(
    ['控制源', '枚举值', '获取锁位置', '释放锁位置'],
    [
        ['触控屏', 'CTL_TOUCH', 'action_star_mixer()', 'action_stop() → fstop()'],
        ['HTTP 网页', 'CTL_HTTP', 'api_start_post_handler()', 'api_stop_post_handler() → fstop()'],
        ['小智 AI', 'CTL_XIAOZHI', 'handle_mcp_command()', 'fstop()'],
    ]
)

add_para('时间开销：互斥锁获取和释放操作均在纳秒级完成（单次 xSemaphoreTake + xSemaphoreGive），不影响 10ms 的 UI 刷新周期。', bold_prefix='技术难点：')

add_heading('2.3.4 DRV8870 PWM 水泵驱动', level=2)

add_para('水泵驱动从 GPIO 继电器开关升级为 DRV8870 电机驱动模块（双路 H 桥），通过 LEDC 硬件外设产生占空比可调的 5kHz PWM 方波，控制水泵电机的转速和启停。', bold_prefix='功能概述：')

add_para('选择 10-bit 分辨率（1024 占空比等级）而非更高位宽的原因是 LEDC 频率可达性约束。ESP32-S3 的 LEDC 时钟源为 80MHz APB 时钟。若采用 13-bit（8192 等级），理论最高频率仅为 9.8kHz，无法达到目标 5kHz，实际输出将远低于预期值。10-bit 模式下最高频率为 78kHz，5kHz 通过分数分频器精确实现。这一约束在初版调试中被忽略，导致电机不转——直到用万用表测得 GPIO12 仅 0.18V 而非预期的 ~1.65V，通过反向计算频率可达性才定位到根因。', bold_prefix='技术难点：')

add_table(
    ['分辨率', '占空比等级', '理论最高频率', '5kHz 是否可达'],
    [
        ['10-bit (LEDC_TIMER_10_BIT)', '0~1023', '78 kHz', '✅ 分数分频器精确实现'],
        ['13-bit (LEDC_TIMER_13_BIT)', '0~8191', '9.8 kHz', '❌ 分频器<1，不可达'],
    ]
)

add_para('LEDC 通道分配需考虑全局唯一性：本项目中 CH0（舵机, Timer2）、CH1（步进, Timer1）、CH2（舵机2, Timer2）、CH6/CH7（ESC×2, Timer3）已被占用，水泵模块采用 CH3（GPIO12, PWM）和 CH4（GPIO13, 低电平），与 Timer0 绑定，8 个通道全部用完。', bold_prefix='资源约束：')

add_heading('2.3.5 小智 AI 智能交互', level=2)

add_para('系统通过微信小程序作为用户入口，集成小智 AI 服务端（xiaozhi-server）实现智能饮食推荐与设备控制。用户在小程序聊天界面中描述当日饮食情况和晚餐面食需求，服务端基于 FastAPI 框架的营养推荐引擎（nutrition-service）解析饮食文本，结合规则引擎与大语言模型识别食物类别、估算摄入量，并根据当日精制主食和杂粮摄入情况，通过推荐算法计算晚餐面团中的杂粮粉配比。推荐结果经用户确认后，服务端通过 WebSocket 长连接以 MCP 协议下发结构化控制指令至 ESP32 设备端，驱动和面机完成全自动制作流程。', bold_prefix='系统概述：')

add_table(
    ['阶段', '小程序/用户端', '服务端处理', '数据/接口'],
    [
        ['饮食输入', '用户文字描述当日饮食和面食需求\n例："中午牛肉面，晚上包300g饺子"', 'LLM 解析 + 规则引擎识别食物条目\n估算每项重量与类别（精制主食/杂粮/蛋白质等）', 'POST /api/v1/intake/parse\n→ ParseResult (FoodItem[] + DoughRequest)'],
        ['确认修正', '用户确认或修改解析出的饮食清单\n调整面团目标重量', '存储确认后的饮食记录至 SQLite\n加载用户画像（年龄组/食欲/消化敏感度）', 'POST /api/v1/recommendations/coarse-grain\n→ Request (confirmed_items + dough_request)'],
        ['营养推荐', '查看推荐结果：杂粮粉克重、普通面粉克重、水量', '推荐算法：根据当日精制主食摄入量（refined_weight）动态计算杂粮比例（10%~40%，步进5%）\n消化风险用户限制杂粮比例≤20%', '→ Recommendation\n(flour_weight_g, water_weight_g, coarse_grain_weight_g, ratio, reason)'],
        ['执行和面', '用户点击"开始和面"', '服务端通过 WebSocket 下发 MCP 指令至 ESP32\nESP32 执行 weight_work(weight)', 'MCP tool: start_mixer\nargs: {weight: 300}'],
    ]
)

add_para('推荐算法的核心公式为：杂粮比例 = 基础 20% +（精制主食 ≥300g 则 +15% / ≥150g 则 +10%）-（已摄入杂粮 ≥100g 则 -5%），最终约束在 10%～40% 范围内并取 5% 整倍数。水量 = 面团总重 × 24%，面粉总量 = 面团总重 - 水量，杂粮粉重 = 面粉总量 × 杂粮比例。例如，用户当日已摄入约 300g 精制主食且无杂粮摄入时，推荐杂粮比例为 35%（20%+15%），300g 饺子的推荐结果为：普通面粉 148g、杂粮粉 80g、水 72g。', bold_prefix='推荐算法：')

add_para('I2S0 配置为全双工模式：BCLK（GPIO14）和 WS（GPIO16）由 INMP441 麦克风与 MAX98357A 扬声器功放共享，SD（GPIO17）接收麦克风数据，DIN（GPIO7）输出扬声器数据。全双工模式仅需 4 个 GPIO，相较于分离式接法节省 2 个引脚。音频采样率 16kHz，单声道，Opus 编码，帧长 60ms。唤醒词检测基于 ESP-SR MultiNet 引擎，关键词"你好小智"，检测后进入录音状态并通过 WebSocket 实时上传音频流。', bold_prefix='音频实现：')

add_para('微信小程序、小智 AI 服务端与 ESP32 设备端的三层架构将用户交互、智能决策和硬件执行分离：小程序承担展示、确认和任务发起的交互职责；服务端承担 NLP 解析、营养推荐计算和任务编排的算法职责；ESP32 承担实时电机控制和安全保护的执行职责。各层职责明确，边界清晰，便于后续独立升级任一模块。', bold_prefix='架构设计原则：')

add_heading('2.3.6 流量计闭环水量控制', level=2)

add_para('将传统定时开环加水（精度 ±20%）升级为 YF-S201 霍尔流量传感器闭环脉冲计数控制（精度 ±2~5%）。传感器输出 450 脉冲/L（约 2.22mL/脉冲），GPIO37 配置为上升沿中断输入，IRAM 中的轻量级 ISR 仅执行 pulse_count++。主循环以 10ms 间隔比较当前脉冲计数与目标值（目标脉冲 = 目标水量 mL × 0.45），达标后立即关断水泵。', bold_prefix='功能概述：')

add_para('在 300g 面团（需水约 73mL, 33 脉冲）场景下，单脉冲量化误差 ±3.0% 叠加传感器 ±2%，总误差约 ±5%。500g 以上面团总误差可控制在 ±4%。水泵关断后管路残余惯性水量约 2~4mL，在传感器固有误差范围内，未采用 PID 补偿以降低复杂度。', bold_prefix='精度分析：')

add_heading('2.3.7 辅助模块', level=2)

add_para('触控屏采用 LVGL v9 图形库，EEZ Studio 可视化设计工具生成页面布局，action.cpp 处理按钮业务逻辑，图形渲染在 lvgl_port_lock 保护下以 10ms 周期执行。HTTP 服务器（端口 80）内嵌网页固件并提供 REST JSON API（/api/start、/api/stop、/api/status），TCP 指令服务器（端口 8080）支持文本行指令，解析采用前缀匹配策略（长前缀优先于短前缀）。两个网络模块在执行电机控制前均需获取互斥锁。')

add_heading('2.3.8 软件技术指标', level=2)

add_table(
    ['指标', '数值', '指标', '数值'],
    [
        ['操作系统', 'FreeRTOS（抢占式）', '音频编码', 'Opus（16kHz/单声道/60ms帧）'],
        ['固件框架', 'ESP-IDF v6.0', '唤醒词引擎', 'ESP-SR MultiNet'],
        ['编程语言', 'C（驱动层）/ C++（应用层）', '服务端模型', 'DeepSeek / ChatGLM'],
        ['图形框架', 'LVGL v9 + EEZ Studio', 'HTTP 服务', '内嵌服务器, 端口 80, REST JSON'],
        ['LEDC 占用', '8/8 通道, 4/4 定时器', 'WiFi 模式', 'AP 模式 (802.11n, 20MHz)'],
        ['固件体积', '2.7MB（工厂分区 3MB, 14%空闲）', 'I2S 模式', '全双工（4 GPIO 双设备）'],
    ]
)

# ── 保存 ──
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '参赛文档_软件系统介绍_v1.docx')
doc.save(out_path)
print(f'✅ 已生成: {out_path}')
